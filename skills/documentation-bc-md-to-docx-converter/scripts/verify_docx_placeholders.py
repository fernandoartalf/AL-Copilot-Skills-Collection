#!/usr/bin/env python3
"""Smoke-test helper: scan a finished ``.docx`` for any surviving ``{{...}}``
placeholders and exit non-zero if any are found.

Usable from the command line::

    python verify_docx_placeholders.py path/to/output.docx

…or imported by any generator's SKILL.md as a smoke test::

    from verify_docx_placeholders import verify
    leftover = verify("output.docx")
    assert not leftover, leftover

The scanner walks the document body, all nested tables, every section's
header and footer (default, first-page, even-page), and returns a sorted
de-duplicated list of remaining tokens (e.g. ``["{{Approvals}}",
"{{Cliente}}"]``).
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError as e:  # pragma: no cover - environment error
    print("[ERROR] python-docx is not installed. Run: pip install python-docx")
    raise SystemExit(1) from e


PLACEHOLDER_PATTERN = re.compile(r"\{\{([^{}]+)\}\}")


def _scan_paragraphs(paragraphs, found: set[str]) -> None:
    for p in paragraphs:
        for m in PLACEHOLDER_PATTERN.finditer(p.text):
            found.add(m.group(0))


def _scan_tables(tables, found: set[str]) -> None:
    for tbl in tables:
        for row in tbl.rows:
            for cell in row.cells:
                _scan_paragraphs(cell.paragraphs, found)
                _scan_tables(cell.tables, found)


def verify(docx_path: str | Path) -> list[str]:
    """Return a sorted list of every ``{{...}}`` token left in *docx_path*.

    An empty list means the document is fully substituted.
    """
    docx_path = Path(docx_path)
    if not docx_path.is_file():
        raise FileNotFoundError(f"docx file not found: {docx_path}")

    doc = Document(str(docx_path))
    found: set[str] = set()

    _scan_paragraphs(doc.paragraphs, found)
    _scan_tables(doc.tables, found)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            try:
                _scan_paragraphs(hf.paragraphs, found)
                _scan_tables(hf.tables, found)
            except Exception:
                pass

    return sorted(found)


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Scan a .docx for unresolved {{Placeholder}} tokens. "
                    "Exits 0 if clean, 1 if any tokens remain."
    )
    parser.add_argument("docx_path", help="Path to the .docx file to verify.")
    parser.add_argument(
        "--quiet",
        action="store_true",
        help="Suppress per-token output; print only the summary count.",
    )
    args = parser.parse_args()

    try:
        leftover = verify(args.docx_path)
    except FileNotFoundError as e:
        print(f"[ERROR] {e}")
        return 1

    if not leftover:
        print(f"[OK] No unresolved placeholders in {args.docx_path}")
        return 0

    if not args.quiet:
        for tok in leftover:
            print(f"[UNRESOLVED] {tok}")
    print(f"[FAIL] {len(leftover)} unresolved placeholder(s) in {args.docx_path}")
    return 1


if __name__ == "__main__":
    sys.exit(main())
