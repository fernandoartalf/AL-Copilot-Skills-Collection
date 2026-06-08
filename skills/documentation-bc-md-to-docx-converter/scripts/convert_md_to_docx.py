"""
Markdown to DOCX Converter using Python

Parses a markdown release note into sections, creates a JSON file with all
section values, then substitutes {{Placeholder}} tokens in a Word (.docx/.dotx)
template with the corresponding JSON values.

Usage:
    python convert_md_to_docx.py <markdown_path> [--template <template_path>] [--output <output_path>] [--open]

Dependencies:
    pip install python-docx
"""

import argparse
import json
import os
import re
import subprocess
import sys
import zipfile
import tempfile
import shutil
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
except ImportError:
    print("[ERROR] python-docx is not installed. Install it with:")
    print("        pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*(\w[\w\s]*?\w|\w)\s*\}\}")
FIELD_MAP_FILENAME = "field-map.json"
UNIFIED_FIELD_MAP_FILENAME = "unified-field-map.json"

# YAML frontmatter delimiter (3 dashes at start of file, then again to close).
FRONTMATTER_PATTERN = re.compile(r"\A---\s*\n(.*?)\n---\s*\n", re.DOTALL)

# Diagram markers emitted by the documentation-bc-ccn-generator skill:
#   <!-- DIAGRAM: mermaid type="<type>" name="<kebab-name>" -->
#   ... optional ASCII fallback ...
#   ```mermaid
#   <actual mermaid source>
#   ```
#   <!-- /DIAGRAM -->
DIAGRAM_MARKER_PATTERN = re.compile(
    r'<!--\s*DIAGRAM:\s*mermaid\s+type="([^"]+)"\s+name="([^"]+)"\s*-->'
    r'(.*?)'
    r'<!--\s*/\s*DIAGRAM\s*-->',
    re.DOTALL | re.IGNORECASE,
)
MERMAID_FENCE_PATTERN = re.compile(r"```mermaid\s*\n(.*?)\n```", re.DOTALL | re.IGNORECASE)
DIAGRAM_TOKEN_PATTERN = re.compile(r"\{\{DIAGRAM:([A-Za-z0-9_\-]+)\}\}")
DIAGRAM_DEFAULT_WIDTH_INCHES = 6.0

# Markdown inline image: ![alt](path "optional title")
# Captures group 1 = alt text, group 2 = path. Title (if present) is ignored.
MD_IMAGE_PATTERN = re.compile(r"!\[([^\]]*)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)")
IMAGE_TOKEN_PATTERN = re.compile(r"\{\{IMAGE:([A-Za-z0-9_\-]+)\}\}")
IMAGE_DEFAULT_WIDTH_INCHES = 6.0
# Combined token splitter used at substitution time.
MEDIA_TOKEN_PATTERN = re.compile(r"\{\{(DIAGRAM|IMAGE):([A-Za-z0-9_\-]+)\}\}")

# Section-key anchor emitted by the documentation-bc-ccn-generator skill
# right after every H2 heading, e.g.:
#   ## 2. Contexto de negocio (procedente de US-001)
#   <!-- section-key: BusinessContext -->
# This decouples the placeholder key from the heading text so the same template
# works regardless of CCN number or content language.
SECTION_KEY_ANCHOR_PATTERN = re.compile(
    r"<!--\s*section-key:\s*([A-Za-z][A-Za-z0-9]*)\s*-->",
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Field Map Loading
# ---------------------------------------------------------------------------

def _shared_dir() -> Path:
    """Return path to .github/skills/shared/."""
    skill_dir = Path(__file__).resolve().parent.parent  # documentation-bc-md-to-docx-converter/
    return skill_dir.parent / "shared"


def _skill_dir() -> Path:
    """Return path to the converter skill root (.github/skills/documentation-bc-md-to-docx-converter/)."""
    return Path(__file__).resolve().parent.parent


def load_field_map() -> dict | None:
    """Deprecated. The legacy local ``field-map.json`` was removed in Phase 1
    of the documentation-bc integration plan; the converter is now driven
    exclusively by ``.github/skills/shared/unified-field-map.json``. Kept as
    a stub so older test scripts that import the symbol don't crash."""
    return None


def resolve_label(field_or_section: dict, lang: str, default_lang: str = "en") -> str:
    """Return a localised label for a registry field or section.

    Reads ``labelLocales[lang]`` (for metadata fields) or
    ``displayHeadingLocales[lang]`` (for sections), with a fallback to
    *default_lang*, then to the canonical ``label`` / ``displayHeading``,
    then to the canonical ``key``.
    """
    locales = (
        field_or_section.get("labelLocales")
        or field_or_section.get("displayHeadingLocales")
        or {}
    )
    if lang in locales and locales[lang]:
        return locales[lang]
    if default_lang in locales and locales[default_lang]:
        return locales[default_lang]
    return (
        field_or_section.get("label")
        or field_or_section.get("displayHeading")
        or field_or_section.get("key")
        or ""
    )


def load_unified_field_map() -> dict | None:
    """Load the unified field map registry, resolving each artifact's source schema.

    The unified map at ``.github/skills/shared/unified-field-map.json`` lists
    every documentation-bc artifact type (USERSTORY, SPEC, ARCHITECTURE,
    ANALYSIS, CCN, PLAN, RELEASENOTE) and points to the canonical
    ``*-fields.json`` owned by its originating skill. This function reads the
    registry, then expands each entry by inlining the referenced schema under
    the ``schema`` key.

    Returns a dict shaped like::

        {
            "artifacts": {
                "USERSTORY": {
                    "artifactType": "USERSTORY",
                    "templateHint": "1_UserStory_Template",
                    "schema": { ...full user-story-fields.json... },
                    ...
                },
                "SPEC":   { ... },
                ...
            }
        }

    or ``None`` if the registry file is missing.
    """
    registry_path = _shared_dir() / UNIFIED_FIELD_MAP_FILENAME
    if not registry_path.exists():
        return None

    with open(registry_path, "r", encoding="utf-8") as f:
        registry = json.load(f)

    resolved: dict[str, dict] = {}
    for entry in registry.get("artifacts", []):
        artifact_type = entry.get("artifactType")
        if not artifact_type:
            continue
        merged = dict(entry)
        source = entry.get("source")
        if source:
            source_path = (registry_path.parent / source).resolve()
            if source_path.exists():
                with open(source_path, "r", encoding="utf-8") as sf:
                    merged["schema"] = json.load(sf)
            else:
                print(f"[WARNING] Schema source not found for {artifact_type}: {source_path}")
                merged["schema"] = None
        else:
            merged["schema"] = None
        resolved[artifact_type] = merged

    return {"artifacts": resolved, "raw": registry}


def parse_yaml_frontmatter(content: str) -> dict[str, str]:
    """Extract YAML frontmatter from a markdown string.

    Returns a flat ``{key: value}`` dict. Performs *light* YAML parsing
    (``key: value`` per line, list values flattened to comma-separated
    strings, nested blocks ignored) so we don't add a PyYAML dependency.
    Returns ``{}`` if no frontmatter block is found.
    """
    match = FRONTMATTER_PATTERN.match(content)
    if not match:
        return {}
    block = match.group(1)
    result: dict[str, str] = {}
    current_key: str | None = None
    list_items: list[str] = []

    def _flush_list():
        nonlocal current_key, list_items
        if current_key is not None and list_items:
            result[current_key] = ", ".join(list_items)
        current_key = None
        list_items = []

    for raw_line in block.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        # List item under previous key
        stripped = line.lstrip()
        if stripped.startswith("- ") and current_key is not None:
            list_items.append(stripped[2:].strip().strip('"\''))
            continue
        # Indented continuation we don't fully parse — skip
        if line.startswith(" ") or line.startswith("\t"):
            continue
        # New top-level key
        _flush_list()
        if ":" in line:
            key, _, value = line.partition(":")
            key = key.strip()
            value = value.strip().strip('"\'')
            if value == "":
                # Could be a list opener
                current_key = key
                list_items = []
                result[key] = ""
            else:
                result[key] = value
                current_key = None
    _flush_list()
    return result


def detect_artifact_type(
    frontmatter: dict[str, str],
    md_path: str,
    registry: dict | None,
) -> str | None:
    """Return the artifactType (USERSTORY / SPEC / ARCHITECTURE / ANALYSIS /
    CCN / PLAN / RELEASENOTE) detected for *md_path*, or ``None``.

    Detection order:
      1. Match the frontmatter ``template:`` value against each artifact's
         registry entry ``template`` / ``templateHint`` (canonical contract
         from each generator's emitted document — see Q1 of
         ``PLAN-documentation-bc-integration.md``).
      2. Match the frontmatter ``id:`` value against each artifact's
         ``schema.idPattern``.
      3. Match the filename against each artifact's ``schema.filePattern``.
    """
    if not registry:
        return None
    artifacts = registry.get("artifacts", {})
    doc_id = frontmatter.get("id", "").strip()
    tpl = (frontmatter.get("template") or "").strip()
    file_name = Path(md_path).name

    # 1. Explicit frontmatter `template:` declaration wins.
    if tpl:
        tpl_stem = Path(tpl).stem
        for art_type, entry in artifacts.items():
            hint = entry.get("templateHint") or ""
            full = entry.get("template") or ""
            if tpl == full or tpl_stem == hint or tpl_stem == Path(full).stem:
                return art_type

    # 2. Frontmatter id matches the artefact's idPattern.
    for art_type, entry in artifacts.items():
        schema = entry.get("schema") or {}
        id_pattern = schema.get("idPattern")
        if id_pattern and doc_id and re.match(id_pattern, doc_id):
            return art_type

    # 3. Filename matches the artefact's filePattern (shell-glob → regex).
    for art_type, entry in artifacts.items():
        schema = entry.get("schema") or {}
        file_pattern = schema.get("filePattern")
        if file_pattern:
            regex = re.escape(file_pattern).replace(r"\*", ".*").replace(r"\?", ".")
            if re.match(regex + r"$", file_name, re.IGNORECASE):
                return art_type

    return None


def _emit_label_tokens(
    schema: dict,
    entry: dict,
    sections: dict[str, str],
    language: str,
    default_lang: str = "en",
) -> None:
    """Populate ``Label_<Key>`` and ``Label_Section_<Key>`` placeholder values
    in *sections* from the artifact schema using *language* (BCP-47).

    Metadata fields are pulled from every list named in ``entry['extraMetadataKeys']``
    plus the standard ``frontmatterFields`` and ``metadataFields`` arrays. Section
    labels come from ``schema['sections']``. *default_lang* is used as the
    fallback when a label is missing in *language*.
    """
    metadata_keys = set(entry.get("extraMetadataKeys") or [])
    metadata_keys.update({"frontmatterFields", "metadataFields"})
    for list_name in metadata_keys:
        for field in schema.get(list_name, []) or []:
            key = field.get("key")
            if not key:
                continue
            sections[f"Label_{key}"] = resolve_label(field, language, default_lang)

    for section in schema.get("sections", []) or []:
        key = section.get("key")
        if not key:
            continue
        sections[f"Label_Section_{key}"] = resolve_label(section, language, default_lang)


# ---------------------------------------------------------------------------
# Markdown Parsing
# ---------------------------------------------------------------------------

def parse_markdown(md_path: str, language: str | None = None, language_fallback: str = "en") -> dict:
    """Parse a markdown document into a dict of named sections.

    Recognises **documentation-bc artifacts** (USERSTORY / SPEC / ARCHITECTURE /
    ANALYSIS / CCN / PLAN / RELEASENOTE) detected via the unified field map
    registry. Detection order: frontmatter ``template:``, then ``id:`` vs
    ``idPattern``, then filename vs ``filePattern``.

    Frontmatter values become placeholders using the ``yamlKey``→``key``
    mapping declared per artifact. Body sections are picked up from
    ``<!-- section-key: ... -->`` anchors emitted right after each H2.

    Localisation: every ``{{Label_<Key>}}`` (metadata) and
    ``{{Label_Section_<Key>}}`` (section heading) placeholder is filled from
    the schema's ``labelLocales`` / ``displayHeadingLocales`` using
    *language*. *language* defaults to the markdown frontmatter ``language:``
    field, then ``"en"``.

    Returns a flat dict like ``{"PlaceholderKey": "Value", ...}``.
    Raises ``ValueError`` if the document cannot be matched to any registered
    artefact type.
    """
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    sections: dict[str, str] = {}

    # --- Detect artifact type via unified field map ---
    registry = load_unified_field_map()
    frontmatter = parse_yaml_frontmatter(content)
    artifact_type = detect_artifact_type(frontmatter, md_path, registry)

    if not artifact_type or not registry:
        raise ValueError(
            f"Could not determine artifact type for '{md_path}'. "
            "Ensure the document has a 'template:' entry in YAML frontmatter "
            "matching one of the registered Word templates "
            "(see .github/skills/shared/unified-field-map.json), or that its "
            "'id:' / filename matches a known idPattern / filePattern."
        )

    entry = registry["artifacts"][artifact_type]
    schema = entry.get("schema") or {}
    resolved_lang = (language or frontmatter.get("language") or "en").strip() or "en"
    print(f"  Artifact type detected: {artifact_type}")
    print(f"  Content language:       {resolved_lang}")

    # Frontmatter → placeholders (yamlKey → key)
    for field in schema.get("frontmatterFields", []):
        key = field.get("key")
        yaml_key = field.get("yamlKey") or key
        if key and yaml_key in frontmatter:
            sections[key] = frontmatter[yaml_key]

    # For SPEC artifacts: emit specMetadataFields PascalCase aliases from the
    # already-resolved camelCase frontmatter keys so templates that use tokens
    # like {{SpecId}}, {{Customer}}, {{ApprovedDate}} etc. also resolve.
    if artifact_type == "SPEC":
        _SPEC_PASCAL_ALIASES: dict[str, str] = {
            "id": "SpecId",
            "title": "SpecTitle",
            "customer": "Customer",
            "status": "SpecStatus",
            "type": "SpecType",
            "userStory": "SourceUserStory",
            "priority": "Priority",
            "complexity": "Complexity",
            "estimatedEffort": "EstimatedEffort",
            "module": "SpecModule",
            "prefix": "Prefix",
            "idRange": "IdRange",
            "createdDate": "CreatedDate",
            "approvedDate": "ApprovedDate",
        }
        for camel_key, pascal_key in _SPEC_PASCAL_ALIASES.items():
            if camel_key in sections and pascal_key not in sections:
                sections[pascal_key] = sections[camel_key]

    # For ARCHITECTURE artifacts: emit archMetadataFields PascalCase aliases
    # from the already-resolved camelCase frontmatter keys so templates that
    # use tokens like {{ArchId}}, {{ArchTitle}}, {{Customer}} etc. also resolve.
    if artifact_type == "ARCHITECTURE":
        _ARCH_PASCAL_ALIASES: dict[str, str] = {
            "id": "ArchId",
            "title": "ArchTitle",
            "customer": "Customer",
            "spec": "SourceSpec",
            "userStory": "SourceUserStory",
            "ccn": "SourceCcn",
            "status": "ArchStatus",
            "version": "ArchVersion",
            "createdDate": "CreatedDate",
            "author": "Author",
        }
        for camel_key, pascal_key in _ARCH_PASCAL_ALIASES.items():
            if camel_key in sections and pascal_key not in sections:
                sections[pascal_key] = sections[camel_key]

    # H1 heading → headingKey placeholder (if declared)
    heading_pattern = entry.get("headingPattern")
    heading_key = entry.get("headingKey")
    if heading_pattern and heading_key:
        h1_match = re.search(
            r"^#\s+" + heading_pattern.lstrip("^"),
            content,
            re.MULTILINE,
        )
        if h1_match:
            sections.setdefault(heading_key, h1_match.group(1).strip())

    # Extra metadata maps (e.g. CCN visible header table)
    for extra_key in entry.get("extraMetadataKeys", []):
        extra_fields = schema.get(extra_key, [])
        label_to_key: dict[str, str] = {}
        for f in extra_fields:
            label = f.get("label")
            key = f.get("key")
            if not (label and key):
                continue
            label_to_key[label] = key
            # Also map locale-specific labels so metadata extraction
            # works when the markdown uses translated labels.
            for locale_label in (f.get("labelLocales") or {}).values():
                if locale_label:
                    label_to_key[locale_label] = key
        for match in re.finditer(
            r"\|\s*\*\*(.+?)\*\*\s*\|\s*(.+?)\s*\|", content
        ):
            lbl = match.group(1).strip()
            val = match.group(2).strip()
            k = label_to_key.get(lbl)
            if k and k not in sections:
                sections[k] = val

    # Emit Label_<Key> / Label_Section_<Key> tokens from the chosen language.
    _emit_label_tokens(schema, entry, sections, resolved_lang, default_lang=language_fallback)
    sections["_language"] = resolved_lang
    sections["_languageFallback"] = language_fallback
    sections["_artifactType"] = artifact_type

    # H2 sections via section-key anchors.
    return _parse_h2_sections(content, sections, md_path=Path(md_path), schema=schema)


def _parse_h2_sections(content: str, sections: dict[str, str], md_path: Path | None = None, schema: dict | None = None) -> dict:
    """Walk every ``## heading`` block in *content* and append each as a
    placeholder value in *sections*. Honours ``<!-- section-key: ... -->``
    anchors as the preferred (language-/id-neutral) key source, falling back
    to a PascalCase form of the heading text otherwise.

    Mermaid diagrams (both ``<!-- DIAGRAM -->`` marker blocks and bare
    triple-backtick ``mermaid`` fences) are extracted into
    ``sections['_diagrams']`` and replaced with ``{{DIAGRAM:<name>}}`` tokens.
    Inline markdown images ``![alt](path)`` are extracted into
    ``sections['_images']`` and replaced with ``{{IMAGE:<name>}}`` tokens.
    """
    h2_splits = re.split(r"^##\s+(.+)$", content, flags=re.MULTILINE)
    # h2_splits: [preamble, heading1, body1, heading2, body2, ...]
    diagrams: dict[str, dict] = {}
    images: dict[str, dict] = {}
    for i in range(1, len(h2_splits) - 1, 2):
        heading = h2_splits[i].strip()
        body = h2_splits[i + 1].strip()

        # Remove leading/trailing horizontal rules (---)
        body = re.sub(r"^---+\s*", "", body)
        body = re.sub(r"\s*---+$", "", body)
        body = body.strip()

        # Prefer an explicit section-key anchor (language-/id-neutral).
        anchor_match = SECTION_KEY_ANCHOR_PATTERN.search(body)
        if anchor_match:
            key = anchor_match.group(1).strip()
            body = SECTION_KEY_ANCHOR_PATTERN.sub("", body, count=1).strip()
        else:
            # Phase-2 safety net: try to resolve the heading against the
            # artefact schema (any registered locale) before falling back to
            # the language-dependent heading_to_key() form. Without this,
            # translated H2s produce translated JSON keys and break the
            # {{Placeholder}} binding in the Word template.
            registry_key = resolve_key_from_schema(heading, schema)
            if registry_key:
                if not getattr(_parse_h2_sections, "_fallback_warned", False):
                    print(
                        f"  [WARN] Section-key anchor missing for H2 '{heading}'. "
                        "Recovered key from schema; please backfill anchors via "
                        ".github/skills/shared/scripts/sync-section-keys.py."
                    )
                    _parse_h2_sections._fallback_warned = True  # type: ignore[attr-defined]
                key = registry_key
            else:
                key = heading_to_key(heading)

        body = extract_diagrams(body, diagrams)
        body = extract_images(body, images, md_path)
        sections[key] = body

    if diagrams:
        sections["_diagrams"] = diagrams
    if images:
        sections["_images"] = images

    return sections


def extract_diagrams(body: str, diagrams: dict[str, dict]) -> str:
    """Find every ``<!-- DIAGRAM: mermaid ... -->...<!-- /DIAGRAM -->`` block in *body*,
    capture its ``name`` / ``type`` / mermaid source, and replace the marker block
    with a ``{{DIAGRAM:<name>}}`` token that the substitution stage will turn into
    an inline image.

    *body* is returned with markers replaced. *diagrams* is mutated in place.
    """

    def _replace(match: re.Match) -> str:
        diagram_type = match.group(1).strip()
        name = match.group(2).strip()
        inner = match.group(3)

        mermaid_match = MERMAID_FENCE_PATTERN.search(inner)
        mermaid_src = mermaid_match.group(1).strip() if mermaid_match else ""

        # Fallback text: the non-mermaid content of the marker block (used when
        # mmdc is not available or the marker has no ```mermaid fenced source).
        fallback = inner
        if mermaid_match:
            fallback = (inner[: mermaid_match.start()] + inner[mermaid_match.end() :]).strip()
        fallback = fallback.strip("`\n ")

        # De-duplicate by name (keep first occurrence; warn on conflict).
        if name in diagrams:
            print(f"[WARNING] Duplicate diagram name '{name}' — keeping first occurrence.")
        else:
            diagrams[name] = {
                "type": diagram_type,
                "name": name,
                "mermaid": mermaid_src,
                "fallback": fallback,
            }
        return f"{{{{DIAGRAM:{name}}}}}"

    body = DIAGRAM_MARKER_PATTERN.sub(_replace, body)

    # Auto-promote any remaining bare ```mermaid``` fences (without surrounding
    # <!-- DIAGRAM --> markers) into synthetic diagram entries.
    def _promote(match: re.Match) -> str:
        mermaid_src = match.group(1).strip()
        # Determine the diagram type from the first non-empty source line.
        first_line = next((ln.strip() for ln in mermaid_src.splitlines() if ln.strip()), "mermaid")
        diagram_type = first_line.split()[0] if first_line else "mermaid"

        existing_auto = sum(1 for n in diagrams if n.startswith("auto-mermaid-"))
        name = f"auto-mermaid-{existing_auto + 1}"
        diagrams[name] = {
            "type": diagram_type,
            "name": name,
            "mermaid": mermaid_src,
            "fallback": "",
            "kind": "diagram",
        }
        return f"{{{{DIAGRAM:{name}}}}}"

    return MERMAID_FENCE_PATTERN.sub(_promote, body)


def extract_images(body: str, images: dict[str, dict], md_path: Path | None) -> str:
    """Find every ``![alt](path)`` markdown image in *body*, resolve its file
    path (relative to the markdown file first, then to the repository root),
    record an entry in *images*, and replace the markdown with an
    ``{{IMAGE:<name>}}`` token. When the file cannot be located the original
    markdown text is left in place and a warning is emitted.
    """
    md_dir = md_path.parent if md_path else None
    # Repo root = four parents up from this script
    # (.github/skills/documentation-bc-md-to-docx-converter/scripts/this.py).
    repo_root = Path(__file__).resolve().parents[4]

    def _resolve(raw_path: str) -> Path | None:
        # Ignore remote URLs (http/https/data:) — those cannot be embedded directly.
        if re.match(r"^[a-zA-Z][a-zA-Z0-9+.\-]*://", raw_path) or raw_path.startswith("data:"):
            return None
        candidates: list[Path] = []
        p = Path(raw_path)
        if p.is_absolute():
            candidates.append(p)
        else:
            if md_dir is not None:
                candidates.append((md_dir / p).resolve())
            candidates.append((repo_root / p).resolve())
        for c in candidates:
            if c.is_file():
                return c
        return None

    def _replace(match: re.Match) -> str:
        alt = match.group(1).strip()
        raw_path = match.group(2).strip()
        resolved = _resolve(raw_path)
        if resolved is None:
            print(f"[WARNING] Image not found, leaving as text: '{raw_path}' (alt='{alt}').")
            return match.group(0)

        existing_auto = sum(1 for n in images if n.startswith("auto-img-"))
        name = f"auto-img-{existing_auto + 1}"
        images[name] = {
            "kind": "image",
            "name": name,
            "alt": alt,
            "src": raw_path,
            "path": str(resolved),
        }
        return f"{{{{IMAGE:{name}}}}}"

    return MD_IMAGE_PATTERN.sub(_replace, body)


def heading_to_key(heading: str) -> str:
    """Convert a heading like 'Change Request Details' to 'ChangeRequestDetails'."""
    # Remove non-alphanumeric chars except spaces
    cleaned = re.sub(r"[^a-zA-Z0-9 ]", "", heading)
    return "".join(word.capitalize() for word in cleaned.split())


def _normalise_heading(text: str) -> str:
    """Normalise an H2 heading for locale-tolerant comparison.

    Strips leading numbering (``1.``, ``1)``, ``1 -``), surrounding whitespace,
    trailing punctuation and lower-cases the result. Used by
    :func:`resolve_key_from_schema` to compare a markdown heading against
    ``displayHeading`` / ``displayHeadingLocales`` entries in the schema.
    """
    if not text:
        return ""
    t = text.strip()
    # Drop a leading section number like "1.", "1)", "1 -", "1\u2013", "1 \u2014"
    t = re.sub(r"^\s*\d+\s*[\.\)\u2013\u2014\-:]?\s*", "", t)
    # Collapse whitespace and casefold
    t = re.sub(r"\s+", " ", t).strip().casefold()
    # Drop trailing punctuation
    t = re.sub(r"[\s\.:;\-\u2013\u2014]+$", "", t)
    return t


def resolve_key_from_schema(heading: str, schema: dict | None) -> str | None:
    """Return the canonical section ``key`` from *schema* whose
    ``displayHeading`` (any locale) matches *heading*, or ``None``.

    Locale-tolerant: normalises both sides via :func:`_normalise_heading`
    (strips numbering, casefolds, trims trailing punctuation) so that a
    translated H2 like ``"1. Referencia a la historia de usuario"`` resolves
    to the canonical key ``"UserStoryReference"`` even when the section-key
    anchor is missing. This is the Phase-2 safety net documented in
    ``docs/plans/PLAN-documentation-bc-section-key-fix.md`` and lets the
    converter recover from legacy/translated files that don't emit anchors.
    """
    if not schema:
        return None
    needle = _normalise_heading(heading)
    if not needle:
        return None
    for section in schema.get("sections", []) or []:
        key = section.get("key")
        if not key:
            continue
        candidates: list[str] = []
        if section.get("displayHeading"):
            candidates.append(section["displayHeading"])
        locales = section.get("displayHeadingLocales") or {}
        candidates.extend(v for v in locales.values() if v)
        for cand in candidates:
            if _normalise_heading(cand) == needle:
                return key
    return None


# ---------------------------------------------------------------------------
# TOC / field-update helpers
# ---------------------------------------------------------------------------

def patch_update_fields(docx_path: str) -> bool:
    """Mark TOC field characters as dirty in word/document.xml so Word
    refreshes the Table of Contents on first open, and remove any global
    <w:updateFields> from word/settings.xml to avoid the "fields that
    may refer to other files" warning dialog.

    Returns True on success, False if the patch was skipped or failed.
    """
    from lxml import etree as lxml_etree

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    SETTINGS_PART = "word/settings.xml"
    DOCUMENT_PART = "word/document.xml"
    UPDATE_TAG = f"{{{W_NS}}}updateFields"
    FLDCHAR_TAG = f"{{{W_NS}}}fldChar"
    INSTRTEXT_TAG = f"{{{W_NS}}}instrText"
    DIRTY_ATTR = f"{{{W_NS}}}dirty"
    FLDCHARTYPE_ATTR = f"{{{W_NS}}}fldCharType"

    tmp_path = docx_path + ".tmp"
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, \
             zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == SETTINGS_PART:
                    # Remove global <w:updateFields> to avoid the
                    # "fields that may refer to other files" dialog.
                    root = lxml_etree.fromstring(data)
                    for existing in root.findall(UPDATE_TAG):
                        root.remove(existing)
                    data = lxml_etree.tostring(
                        root,
                        xml_declaration=True,
                        encoding="UTF-8",
                        standalone=True,
                    )

                elif item.filename == DOCUMENT_PART:
                    # Walk field characters and mark TOC fields as dirty.
                    root = lxml_etree.fromstring(data)
                    in_toc = False
                    for fc in root.iter(FLDCHAR_TAG):
                        fc_type = fc.get(FLDCHARTYPE_ATTR, "")
                        if fc_type == "begin":
                            run = fc.getparent()
                            if run is not None:
                                paragraph = run.getparent()
                                if paragraph is not None:
                                    found_toc = False
                                    started = False
                                    for sibling in paragraph:
                                        if sibling is run:
                                            started = True
                                            continue
                                        if not started:
                                            continue
                                        for child in sibling:
                                            if child.tag == INSTRTEXT_TAG and child.text and "TOC" in child.text:
                                                found_toc = True
                                                break
                                            if child.tag == FLDCHAR_TAG:
                                                break
                                        if found_toc:
                                            break
                                    if found_toc:
                                        in_toc = True
                                        fc.set(DIRTY_ATTR, "true")
                        elif fc_type == "end" and in_toc:
                            in_toc = False
                    data = lxml_etree.tostring(
                        root,
                        xml_declaration=True,
                        encoding="UTF-8",
                        standalone=True,
                    )

                zout.writestr(item, data)
        shutil.move(tmp_path, docx_path)
        return True
    except Exception as exc:
        if os.path.isfile(tmp_path):
            os.remove(tmp_path)
        print(f"[WARNING] Could not patch TOC dirty flags: {exc}")
        return False


# ---------------------------------------------------------------------------
# JSON helpers
# ---------------------------------------------------------------------------

def write_json(sections: dict, json_path: str) -> str:
    """Write sections dict to a JSON file and return the path."""
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(sections, f, indent=2, ensure_ascii=False)
    return json_path


# ---------------------------------------------------------------------------
# DOCX Template Substitution
# ---------------------------------------------------------------------------

# Cache of rendered mermaid sources → PNG paths, so the same diagram
# referenced from multiple placeholders (or via repeated tokens) is only
# rendered to disk once per conversion run.
_MERMAID_RENDER_CACHE: dict[str, Path | None] = {}
_MMDC_AVAILABILITY: bool | None = None


def _mmdc_available() -> bool:
    """Return True if the Mermaid CLI (``mmdc``) is on PATH."""
    global _MMDC_AVAILABILITY
    if _MMDC_AVAILABILITY is None:
        _MMDC_AVAILABILITY = shutil.which("mmdc") is not None
        if not _MMDC_AVAILABILITY:
            print(
                "[WARNING] Mermaid CLI (mmdc) not found on PATH. Diagrams will be "
                "inserted as their ASCII fallback text. Install with: "
                "`npm install -g @mermaid-js/mermaid-cli`."
            )
    return _MMDC_AVAILABILITY


def render_mermaid_to_png(name: str, mermaid_src: str) -> Path | None:
    """Render a Mermaid source string to a PNG file and return its path.

    Returns ``None`` if ``mmdc`` is not installed, the source is empty, or
    rendering fails. The result is cached per-run keyed by the source text.
    """
    if not mermaid_src.strip():
        return None
    cache_key = mermaid_src
    if cache_key in _MERMAID_RENDER_CACHE:
        return _MERMAID_RENDER_CACHE[cache_key]
    if not _mmdc_available():
        _MERMAID_RENDER_CACHE[cache_key] = None
        return None

    tmpdir = Path(tempfile.gettempdir()) / "md-to-docx-mermaid"
    tmpdir.mkdir(exist_ok=True)
    safe_name = re.sub(r"[^A-Za-z0-9_\-]", "_", name) or "diagram"
    mmd_file = tmpdir / f"{safe_name}.mmd"
    png_file = tmpdir / f"{safe_name}.png"
    mmd_file.write_text(mermaid_src, encoding="utf-8")

    try:
        subprocess.run(
            [
                shutil.which("mmdc") or "mmdc",
                "-i", str(mmd_file),
                "-o", str(png_file),
                "-b", "white",
                "--quiet",
            ],
            check=True,
            capture_output=True,
        )
    except (FileNotFoundError, subprocess.CalledProcessError) as exc:
        stderr = getattr(exc, "stderr", b"") or b""
        if isinstance(stderr, bytes):
            stderr = stderr.decode("utf-8", errors="ignore")
        print(f"[WARNING] Failed to render mermaid diagram '{name}': {exc}\n{stderr}")
        _MERMAID_RENDER_CACHE[cache_key] = None
        return None

    if not png_file.exists():
        _MERMAID_RENDER_CACHE[cache_key] = None
        return None

    _MERMAID_RENDER_CACHE[cache_key] = png_file
    return png_file


def process_markdown_formatting(text: str) -> str:
    """Process markdown in text:
    - Remove all asterisk (*) symbols used for bold/italic formatting
    """
    # Remove all asterisks (used for bold ** and italic * in markdown)
    text = text.replace('*', '')
    
    return text


def copy_run_format(source_run, target_run):
    """Copy all formatting properties from source run to target run.
    
    Preserves: font name, size, color, bold, italic, underline, highlight, etc.
    """
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb
    if source_run.font.bold is not None:
        target_run.font.bold = source_run.font.bold
    if source_run.font.italic is not None:
        target_run.font.italic = source_run.font.italic
    if source_run.font.underline is not None:
        target_run.font.underline = source_run.font.underline
    if source_run.font.strike is not None:
        target_run.font.strike = source_run.font.strike
    if source_run.font.all_caps is not None:
        target_run.font.all_caps = source_run.font.all_caps
    if source_run.font.small_caps is not None:
        target_run.font.small_caps = source_run.font.small_caps
    if source_run.font.highlight_color:
        target_run.font.highlight_color = source_run.font.highlight_color


def apply_formatted_text_to_paragraph(paragraph, text: str, template_run=None, diagrams: dict | None = None):
    """Apply text to paragraph with markdown formatting while preserving template formatting.

    Args:
        paragraph: The paragraph to update
        text: The replacement text
        template_run: The original run from template to copy formatting from
        diagrams: Optional dict mapping ``name`` -> media entry produced by
            :func:`extract_diagrams` and :func:`extract_images`. Both
            ``{{DIAGRAM:<name>}}`` and ``{{IMAGE:<name>}}`` tokens are
            looked up in this dict. Diagram entries are rendered via
            ``mmdc``; image entries are embedded directly from disk.

    - Lines starting with # become bold
    - Template formatting (font, size, color) is preserved
    - ``{{DIAGRAM:<name>}}`` / ``{{IMAGE:<name>}}`` tokens are turned into inline images
    """
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""

    media = diagrams or {}

    # Split text by media tokens. ``re.split`` with two capture groups yields:
    # [text_before, kind_1, name_1, text_between, kind_2, name_2, ...]
    parts = MEDIA_TOKEN_PATTERN.split(text)

    i = 0
    while i < len(parts):
        # Text segment.
        text_part = parts[i]
        if text_part:
            _add_formatted_text_segment(paragraph, text_part, template_run)
        i += 1
        if i >= len(parts):
            break
        # Captured token: kind, name.
        kind = parts[i]
        name = parts[i + 1]
        i += 2
        entry = media.get(name)
        if entry is None:
            print(f"[WARNING] {kind} token '{name}' has no matching entry in the JSON.")
            _add_formatted_text_segment(paragraph, f"[{kind.title()}: {name}]", template_run)
            continue

        if kind == "IMAGE" or entry.get("kind") == "image":
            _embed_image(paragraph, entry, template_run)
        else:
            _embed_diagram(paragraph, entry, template_run)


def _embed_diagram(paragraph, diagram: dict, template_run):
    """Render a mermaid diagram entry to PNG and embed it inline."""
    name = diagram.get("name", "?")
    png_path = render_mermaid_to_png(name, diagram.get("mermaid", ""))
    if png_path is not None:
        image_run = paragraph.add_run()
        try:
            image_run.add_picture(str(png_path), width=Inches(DIAGRAM_DEFAULT_WIDTH_INCHES))
            return
        except Exception as exc:  # pragma: no cover - defensive
            print(f"[WARNING] Could not embed image for diagram '{name}': {exc}")
    _add_diagram_fallback(paragraph, diagram, template_run)


def _embed_image(paragraph, image: dict, template_run):
    """Embed an inline ``![alt](path)`` image into the paragraph."""
    name = image.get("name", "?")
    path = image.get("path")
    alt = image.get("alt") or name
    if not path or not Path(path).is_file():
        print(f"[WARNING] Image file missing for '{name}': {path}")
        _add_formatted_text_segment(paragraph, f"[Image: {alt} (file not found)]", template_run)
        return
    image_run = paragraph.add_run()
    try:
        image_run.add_picture(str(path), width=Inches(IMAGE_DEFAULT_WIDTH_INCHES))
    except Exception as exc:  # pragma: no cover - defensive
        print(f"[WARNING] Could not embed image '{name}' ({path}): {exc}")
        _add_formatted_text_segment(paragraph, f"[Image: {alt}]", template_run)


def _add_formatted_text_segment(paragraph, text: str, template_run):
    """Append *text* to *paragraph* honouring the legacy heading-bold behaviour."""
    lines = text.split('\n')
    first = True
    for line in lines:
        is_heading = line.lstrip().startswith('#')
        if is_heading:
            clean_line = re.sub(r'^#+\s*', '', line.lstrip())
            run = paragraph.add_run(clean_line if first else '\n' + clean_line)
            if template_run:
                copy_run_format(template_run, run)
            run.bold = True
        else:
            run = paragraph.add_run(line if first else '\n' + line)
            if template_run:
                copy_run_format(template_run, run)
        first = False


def _add_diagram_fallback(paragraph, diagram: dict, template_run):
    """Insert the diagram's plain-text fallback when no PNG can be rendered."""
    fallback = diagram.get("fallback") or diagram.get("mermaid") or ""
    label = f"[Diagram: {diagram.get('name', '?')} ({diagram.get('type', 'mermaid')})]"
    text = label if not fallback else f"{label}\n{fallback}"
    _add_formatted_text_segment(paragraph, text, template_run)


# ---------------------------------------------------------------------------
# Markdown mini-renderer (Phase 5)
#
# When a {{Key}} placeholder resolves to a value containing block-level
# markdown (GFM tables, lists, fenced code, headings, multi-paragraph), we
# parse the value into typed blocks and emit native Word constructs at the
# placeholder paragraph's location. The placeholder paragraph is then
# removed. Media tokens ({{DIAGRAM:..}} / {{IMAGE:..}}) are preserved through
# the parser and embedded inline by ``_apply_inline_md``.
# ---------------------------------------------------------------------------

_MD_TABLE_LINE_RE = re.compile(r"^\s*\|.*\|\s*$")
_MD_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
_MD_BULLET_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
_MD_ORDERED_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")
_MD_FENCE_RE = re.compile(r"^\s*```\s*([\w+-]*)\s*$")
_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*#*\s*$")

# Inline tokens. Order: bold, code, link, italic_star, italic_underscore.
_MD_INLINE_RE = re.compile(
    r"(\*\*[^*\n]+?\*\*)"
    r"|(`[^`\n]+?`)"
    r"|(\[[^\]\n]+?\]\([^)\n]+?\))"
    r"|((?<![\w*])\*[^*\n]+?\*(?!\w))"
    r"|((?<![\w_])_[^_\n]+?_(?!\w))"
)
_MD_INLINE_LINK_RE = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")


def _looks_like_block_markdown(text: str) -> bool:
    """Return True when *text* should be rendered via the block parser."""
    if not text:
        return False
    if "```" in text:
        return True
    has_newline = "\n" in text
    for line in text.splitlines():
        if (_MD_TABLE_LINE_RE.match(line)
                or _MD_BULLET_RE.match(line)
                or _MD_ORDERED_RE.match(line)
                or _MD_HEADING_RE.match(line)):
            return True
    if has_newline and "\n\n" in text.strip():
        return True
    return False


def _split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def parse_markdown_blocks(text: str) -> list[dict]:
    """Tokenize *text* into a list of typed block dicts.

    Block types: ``paragraph``, ``heading``, ``bullet``, ``ordered``,
    ``table``, ``code``. Inline markdown is kept verbatim inside the block
    and resolved later by :func:`_apply_inline_md`.
    """
    # Normalise CRLF.
    lines = text.replace("\r\n", "\n").split("\n")
    blocks: list[dict] = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # Fenced code block.
        fm = _MD_FENCE_RE.match(line)
        if fm:
            lang = fm.group(1) or ""
            buf: list[str] = []
            i += 1
            while i < n and not _MD_FENCE_RE.match(lines[i]):
                buf.append(lines[i])
                i += 1
            if i < n:
                i += 1  # consume closing fence
            blocks.append({"type": "code", "language": lang, "text": "\n".join(buf)})
            continue

        # Heading.
        hm = _MD_HEADING_RE.match(line)
        if hm:
            blocks.append({
                "type": "heading",
                "level": len(hm.group(1)),
                "text": hm.group(2).strip(),
            })
            i += 1
            continue

        # GFM table: header row followed by separator row.
        if (_MD_TABLE_LINE_RE.match(line)
                and i + 1 < n
                and _MD_TABLE_SEP_RE.match(lines[i + 1])):
            header = _split_table_row(line)
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < n and _MD_TABLE_LINE_RE.match(lines[i]):
                rows.append(_split_table_row(lines[i]))
                i += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue

        # Bullet list.
        if _MD_BULLET_RE.match(line):
            items: list[str] = []
            while i < n and _MD_BULLET_RE.match(lines[i]):
                items.append(_MD_BULLET_RE.match(lines[i]).group(2))
                i += 1
            blocks.append({"type": "bullet", "items": items})
            continue

        # Ordered list.
        if _MD_ORDERED_RE.match(line):
            items = []
            while i < n and _MD_ORDERED_RE.match(lines[i]):
                items.append(_MD_ORDERED_RE.match(lines[i]).group(2))
                i += 1
            blocks.append({"type": "ordered", "items": items})
            continue

        # Blank line.
        if not line.strip():
            i += 1
            continue

        # Plain paragraph: gather until blank line or another block construct.
        para: list[str] = []
        while i < n:
            cur = lines[i]
            if not cur.strip():
                break
            if (_MD_FENCE_RE.match(cur)
                    or _MD_HEADING_RE.match(cur)
                    or _MD_BULLET_RE.match(cur)
                    or _MD_ORDERED_RE.match(cur)
                    or (_MD_TABLE_LINE_RE.match(cur)
                        and i + 1 < n
                        and _MD_TABLE_SEP_RE.match(lines[i + 1]))):
                break
            para.append(cur.strip())
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para)})
    return blocks


def _apply_inline_md(paragraph, text: str, template_run, media: dict | None):
    """Append runs to *paragraph* applying inline markdown + media tokens.

    Handles ``**bold**``, ``*italic*`` / ``_italic_``, `` `code` ``,
    ``[label](url)`` and inline ``{{DIAGRAM:..}}`` / ``{{IMAGE:..}}``.
    Existing runs in *paragraph* are NOT cleared.
    """
    if not text:
        return
    # Split by media tokens first so embedded images survive intact.
    parts = MEDIA_TOKEN_PATTERN.split(text)
    i = 0
    while i < len(parts):
        chunk = parts[i]
        if chunk:
            _append_inline_runs(paragraph, chunk, template_run)
        i += 1
        if i >= len(parts):
            break
        kind = parts[i]
        name = parts[i + 1]
        i += 2
        entry = (media or {}).get(name)
        if entry is None:
            r = paragraph.add_run(f"[{kind.title()}: {name}]")
            if template_run:
                copy_run_format(template_run, r)
            continue
        if kind == "IMAGE" or entry.get("kind") == "image":
            _embed_image(paragraph, entry, template_run)
        else:
            _embed_diagram(paragraph, entry, template_run)


def _append_inline_runs(paragraph, text: str, template_run):
    """Tokenize *text* and append styled runs to *paragraph*."""
    pos = 0
    for m in _MD_INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            if template_run:
                copy_run_format(template_run, run)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            if template_run:
                copy_run_format(template_run, run)
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            if template_run:
                copy_run_format(template_run, run)
            run.font.name = "Consolas"
        elif token.startswith("["):
            lm = _MD_INLINE_LINK_RE.match(token)
            if lm:
                run = paragraph.add_run(lm.group(1))
                if template_run:
                    copy_run_format(template_run, run)
                run.underline = True
                try:
                    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                except Exception:
                    pass
            else:
                run = paragraph.add_run(token)
                if template_run:
                    copy_run_format(template_run, run)
        elif (token.startswith("*") and token.endswith("*")) or (token.startswith("_") and token.endswith("_")):
            run = paragraph.add_run(token[1:-1])
            if template_run:
                copy_run_format(template_run, run)
            run.italic = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if template_run:
            copy_run_format(template_run, run)


def _clone_paragraph_skeleton(template_p_elem):
    """Clone the placeholder paragraph element, keeping ``w:pPr`` and dropping
    all runs/hyperlinks so callers can append fresh runs."""
    from copy import deepcopy
    new_p = deepcopy(template_p_elem)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    return new_p


def _new_block_paragraph(template_p_elem, paragraph_parent):
    """Create a sibling paragraph (not yet inserted) cloned from the
    placeholder paragraph's pPr, returning the lxml element and a wrapping
    :class:`Paragraph`."""
    new_p = _clone_paragraph_skeleton(template_p_elem)
    return new_p, Paragraph(new_p, paragraph_parent)


def _build_table_element(header: list[str], rows: list[list[str]], template_run, media: dict | None):
    """Construct a ``<w:tbl>`` lxml element from a parsed GFM table block."""
    n_cols = max(len(header), max((len(r) for r in rows), default=0), 1)
    header = list(header) + [""] * (n_cols - len(header))
    rows = [list(r) + [""] * (n_cols - len(r)) for r in rows]

    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    for _ in range(n_cols):
        tblGrid.append(OxmlElement("w:gridCol"))
    tbl.append(tblGrid)

    def _add_row(cells, bold):
        tr = OxmlElement("w:tr")
        for cell_text in cells:
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), "0")
            tcW.set(qn("w:type"), "auto")
            tcPr.append(tcW)
            tc.append(tcPr)
            p = OxmlElement("w:p")
            tc.append(p)
            para = Paragraph(p, None)
            _apply_inline_md(para, cell_text, template_run, media)
            if bold:
                for run in para.runs:
                    run.bold = True
            tr.append(tc)
        tbl.append(tr)

    _add_row(header, bold=True)
    for row in rows:
        _add_row(row, bold=False)
    return tbl


def render_markdown_blocks_into(paragraph, blocks: list[dict], template_run, media: dict | None):
    """Insert each block as a sibling element BEFORE *paragraph*, then remove
    *paragraph* from its parent.

    Heading blocks render bold. Bullet/ordered lists fall back to literal
    bullet prefixes when the document does not define list styles. Tables
    are emitted as native Word tables with the placeholder paragraph's
    surrounding context. Code blocks render monospace.
    """
    p_elem = paragraph._element
    parent = p_elem.getparent()
    if parent is None:
        # Not attached to a tree; fall back to inline rendering.
        _apply_inline_md(paragraph, "\n".join(_block_to_text(b) for b in blocks), template_run, media)
        return

    paragraph_parent = paragraph._parent

    for block in blocks:
        btype = block["type"]
        if btype == "paragraph":
            new_p, new_para = _new_block_paragraph(p_elem, paragraph_parent)
            p_elem.addprevious(new_p)
            _apply_inline_md(new_para, block["text"], template_run, media)
        elif btype == "heading":
            new_p, new_para = _new_block_paragraph(p_elem, paragraph_parent)
            p_elem.addprevious(new_p)
            _apply_inline_md(new_para, block["text"], template_run, media)
            for run in new_para.runs:
                run.bold = True
        elif btype == "bullet":
            for item in block["items"]:
                new_p, new_para = _new_block_paragraph(p_elem, paragraph_parent)
                p_elem.addprevious(new_p)
                prefix = new_para.add_run("• ")
                if template_run:
                    copy_run_format(template_run, prefix)
                _apply_inline_md(new_para, item, template_run, media)
        elif btype == "ordered":
            for idx, item in enumerate(block["items"], 1):
                new_p, new_para = _new_block_paragraph(p_elem, paragraph_parent)
                p_elem.addprevious(new_p)
                prefix = new_para.add_run(f"{idx}. ")
                if template_run:
                    copy_run_format(template_run, prefix)
                _apply_inline_md(new_para, item, template_run, media)
        elif btype == "code":
            for code_line in block["text"].split("\n"):
                new_p, new_para = _new_block_paragraph(p_elem, paragraph_parent)
                p_elem.addprevious(new_p)
                run = new_para.add_run(code_line)
                if template_run:
                    copy_run_format(template_run, run)
                run.font.name = "Consolas"
        elif btype == "table":
            tbl = _build_table_element(block["header"], block["rows"], template_run, media)
            p_elem.addprevious(tbl)

    parent.remove(p_elem)


def _block_to_text(block: dict) -> str:
    """Flatten a parsed block back to plain text (used only as a safety
    fallback when a paragraph is detached from a parent tree)."""
    btype = block["type"]
    if btype in ("paragraph", "heading"):
        return block.get("text", "")
    if btype in ("bullet", "ordered"):
        return "\n".join(block.get("items", []))
    if btype == "code":
        return block.get("text", "")
    if btype == "table":
        out = [" | ".join(block.get("header", []))]
        for row in block.get("rows", []):
            out.append(" | ".join(row))
        return "\n".join(out)
    return ""


# ---------------------------------------------------------------------------


def replace_placeholder_in_paragraph(paragraph, sections: dict, diagrams: dict | None = None) -> bool:
    """Replace {{Key}} placeholders in a paragraph's runs.

    Handles placeholders split across multiple runs by joining all run texts,
    performing the replacement, then rewriting the runs.
    Preserves formatting (font, size, color, bold, etc.) from the template
    placeholder. When the substituted value contains block-level markdown
    (GFM tables, lists, fenced code, headings, multi-paragraph) the value
    is parsed into typed blocks and emitted as native Word constructs at
    the placeholder's location via :func:`render_markdown_blocks_into`; the
    placeholder paragraph itself is then removed. Otherwise the legacy
    single-paragraph inline rendering path is used.

    Returns True if any replacement was made.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not PLACEHOLDER_PATTERN.search(full_text):
        return False

    def _replacer(m):
        key = m.group(1).strip()
        return sections.get(key, m.group(0))  # keep original if key not found

    new_text = PLACEHOLDER_PATTERN.sub(_replacer, full_text)
    if new_text == full_text:
        return False

    # Capture formatting from the first run that has content (template formatting).
    template_run = None
    for run in paragraph.runs:
        if run.text.strip():
            template_run = run
            break
    if not template_run and paragraph.runs:
        template_run = paragraph.runs[0]

    # Phase 5: rich block-level markdown path.
    # Skip for Heading styles: the template embeds manual "N. " prefixes in
    # heading paragraphs which would be misdetected as markdown ordered-list
    # markers, causing every section heading to render as "1.".
    try:
        para_style_name = paragraph.style.name if paragraph.style else ""
    except AttributeError:
        para_style_name = ""
    is_heading_para = para_style_name.lower().startswith("heading")
    if not is_heading_para and _looks_like_block_markdown(new_text):
        blocks = parse_markdown_blocks(new_text)
        if blocks:
            # Clear placeholder paragraph runs so any surrounding chrome
            # (which we just consumed via full_text) doesn't linger.
            for run in paragraph.runs:
                run.text = ""
            render_markdown_blocks_into(paragraph, blocks, template_run, diagrams)
            return True

    # Legacy single-paragraph path (strip asterisks, render inline).
    new_text = process_markdown_formatting(new_text)
    apply_formatted_text_to_paragraph(paragraph, new_text, template_run, diagrams=diagrams)
    return True



def replace_placeholder_in_table_cell(cell, sections: dict, diagrams: dict | None = None):
    """Replace placeholders in every paragraph of a table cell."""
    for paragraph in cell.paragraphs:
        replace_placeholder_in_paragraph(paragraph, sections, diagrams=diagrams)
    # Cells can contain nested tables
    for nested_table in cell.tables:
        replace_placeholders_in_table(nested_table, sections, diagrams=diagrams)


def replace_placeholders_in_table(table, sections: dict, diagrams: dict | None = None):
    """Replace placeholders across all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            replace_placeholder_in_table_cell(cell, sections, diagrams=diagrams)


def replace_placeholders_in_textboxes(container_element, sections: dict, diagrams: dict | None = None):
    """Replace {{Key}} placeholders inside all textboxes within a container XML element.

    Textboxes (floating or inline) are stored as ``w:txbxContent`` elements inside
    ``w:drawing`` shapes and are invisible to python-docx's ``.paragraphs`` property.
    This function walks the raw lxml tree, finds every ``w:txbxContent``, and
    processes each contained ``w:p`` as a regular paragraph.
    """
    for txbx_content in container_element.iter(qn("w:txbxContent")):
        for p_elem in txbx_content.findall(qn("w:p")):
            paragraph = Paragraph(p_elem, None)
            replace_placeholder_in_paragraph(paragraph, sections, diagrams=diagrams)


def replace_placeholders_in_document(doc: Document, sections: dict, diagrams: dict | None = None):
    """Walk every paragraph, table, header, and footer in the document
    and replace {{Key}} placeholders with values from *sections*.
    Inline ``{{DIAGRAM:<name>}}`` tokens are rendered as images using *diagrams*."""

    # Body paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholder_in_paragraph(paragraph, sections, diagrams=diagrams)

    # Body tables
    for table in doc.tables:
        replace_placeholders_in_table(table, sections, diagrams=diagrams)

    # Body textboxes (floating shapes not exposed by doc.paragraphs)
    replace_placeholders_in_textboxes(doc.element.body, sections, diagrams=diagrams)

    # Headers & footers in all sections
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header:
                for paragraph in header.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, sections, diagrams=diagrams)
                for table in header.tables:
                    replace_placeholders_in_table(table, sections, diagrams=diagrams)
                # Textboxes inside headers
                replace_placeholders_in_textboxes(header._element, sections, diagrams=diagrams)
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer:
                for paragraph in footer.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, sections, diagrams=diagrams)
                for table in footer.tables:
                    replace_placeholders_in_table(table, sections, diagrams=diagrams)
                # Textboxes inside footers
                replace_placeholders_in_textboxes(footer._element, sections, diagrams=diagrams)


# ---------------------------------------------------------------------------
# DOTX Template Handling
# ---------------------------------------------------------------------------

def open_template_document(template_path: str) -> Document:
    """Open a .dotx or .docx file, bypassing strict content type checks.
    
    python-docx >=1.1.0 rejects .dotx files with template content types.
    This function works around that by using zipfile to extract and reload.
    """
    try:
        # First try the normal way
        doc = Document(template_path)
        return doc
    except ValueError as e:
        # If the error is about content type, use a workaround
        if 'is not a Word file' in str(e) or 'content type' in str(e):
            # Create a temporary copy and modify the content type
            with tempfile.TemporaryDirectory() as tmpdir:
                # Copy the template to a temp location
                temp_template = os.path.join(tmpdir, "template_copy.dotx")
                shutil.copy2(template_path, temp_template)
                
                # Open as a zip and modify [Content_Types].xml
                with zipfile.ZipFile(temp_template, 'r') as zip_read:
                    # Extract all files
                    extracted_dir = os.path.join(tmpdir, "extracted")
                    zip_read.extractall(extracted_dir)
                
                # Modify the content type in [Content_Types].xml
                content_types_path = os.path.join(extracted_dir, "[Content_Types].xml")
                if os.path.exists(content_types_path):
                    with open(content_types_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    # Replace template content type with document content type
                    content = content.replace(
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml',
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
                    )
                    
                    with open(content_types_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                
                # Re-zip as a .docx
                temp_docx = os.path.join(tmpdir, "template.docx")
                with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zip_write:
                    for root, dirs, files in os.walk(extracted_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arc_name = os.path.relpath(file_path, extracted_dir)
                            zip_write.write(file_path, arc_name)
                
                # Now open the modified document
                doc = Document(temp_docx)
                return doc
        else:
            raise


# ---------------------------------------------------------------------------
# Template discovery
# ---------------------------------------------------------------------------

def get_examples_dir() -> Path:
    """Return the path to the skill's templates/ folder."""
    return Path(__file__).resolve().parent.parent / "templates"


def list_templates() -> list[Path]:
    """Return all .docx and .dotx template files in the examples/ folder."""
    examples = get_examples_dir()
    if not examples.is_dir():
        return []
    templates = sorted(
        p for p in examples.iterdir()
        if p.suffix.lower() in (".docx", ".dotx")
        and not p.name.startswith("~$")
    )
    return templates


def find_template(start_dir: str) -> str | None:
    """Look for a template file.

    If only one template exists in the examples folder, return it.
    If multiple templates exist, return None (caller should prompt the user).
    If no templates exist in examples, walk up from start_dir as a fallback.
    """
    templates = list_templates()
    if len(templates) == 1:
        return str(templates[0])
    if len(templates) > 1:
        # Multiple templates — caller must prompt
        return None

    return None


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def _scan_placeholders(doc) -> list[str]:
    """Return a sorted, de-duplicated list of every ``{{...}}`` token present
    anywhere in the document body, nested tables, headers, or footers.
    Used both before substitution (to inventory the template) and after
    substitution (to detect unresolved tokens)."""
    found: set[str] = set()

    def _scan_paragraphs(paragraphs):
        for p in paragraphs:
            for m in PLACEHOLDER_PATTERN.finditer(p.text):
                found.add(m.group(0))

    def _scan_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    _scan_paragraphs(cell.paragraphs)
                    _scan_tables(cell.tables)

    _scan_paragraphs(doc.paragraphs)
    _scan_tables(doc.tables)
    for section in doc.sections:
        for hf in (section.header, section.footer, section.first_page_header,
                   section.first_page_footer, section.even_page_header,
                   section.even_page_footer):
            try:
                _scan_paragraphs(hf.paragraphs)
                _scan_tables(hf.tables)
            except Exception:
                pass

    return sorted(found)


def _scan_unresolved_placeholders(doc) -> list[str]:
    """Backward-compatible alias of :func:`_scan_placeholders` used by the
    strict-mode check after substitution."""
    return _scan_placeholders(doc)


def convert(
    md_path: str,
    template_path: str,
    output_path: str,
    open_after: bool = False,
    language: str | None = None,
    language_fallback: str = "en",
    strict: bool = True,
    report: bool = False,
):
    """Full pipeline: MD → JSON → DOCX."""

    md_path = os.path.abspath(md_path)
    template_path = os.path.abspath(template_path)
    output_path = os.path.abspath(output_path)

    print()
    print("=" * 56)
    print(" Markdown to DOCX Converter (Python + python-docx)")
    print("=" * 56)
    print()

    # 1. Validate inputs
    if not os.path.isfile(md_path):
        print(f"[ERROR] Markdown file not found: {md_path}")
        sys.exit(1)
    print(f"[OK] Markdown file : {md_path}")

    if not os.path.isfile(template_path):
        print(f"[ERROR] Template file not found: {template_path}")
        sys.exit(1)
    print(f"[OK] Template file : {template_path}")

    # 1b. Report unified field map status
    registry = load_unified_field_map()
    if registry:
        n_art = len(registry.get("artifacts", {}))
        print(f"[OK] Unified map   : {n_art} artifact type(s) registered")
    else:
        print("[WARNING] Unified field map not found at .github/skills/shared/unified-field-map.json")

    # 2. Parse markdown into sections
    print()
    print("Parsing markdown...")
    sections = parse_markdown(md_path, language=language, language_fallback=language_fallback)
    diagrams = sections.pop("_diagrams", {}) if isinstance(sections, dict) else {}
    images = sections.pop("_images", {}) if isinstance(sections, dict) else {}
    # Merge images into the media dict threaded as ``diagrams=`` through the
    # substitution layer. Keys are namespace-prefixed (``auto-img-*`` vs
    # ``auto-mermaid-*``/explicit diagram names) so collisions are not expected.
    media = {**diagrams, **images}
    artifact_type = sections.get("_artifactType")
    resolved_lang = sections.get("_language")
    print(f"  Found {len(sections)} placeholder value(s)")
    if diagrams:
        print(f"  Found {len(diagrams)} diagram(s): {', '.join(diagrams.keys())}")
    if images:
        print(f"  Found {len(images)} image(s): {', '.join(images.keys())}")

    # 3. Write JSON file alongside the output
    json_path = os.path.splitext(output_path)[0] + ".json"
    json_payload = dict(sections)
    if diagrams:
        json_payload["_diagrams"] = diagrams
    if images:
        json_payload["_images"] = images
    write_json(json_payload, json_path)
    print(f"[OK] JSON file     : {json_path}")

    # 4. Open template and perform substitution
    print()
    print("Applying template substitutions...")
    doc = open_template_document(template_path)
    template_tokens = _scan_placeholders(doc)
    replace_placeholders_in_document(doc, sections, diagrams=media)

    # 4b. Strict-mode check: any remaining {{...}} tokens are errors.
    unresolved = _scan_unresolved_placeholders(doc)
    substituted = sorted(set(template_tokens) - set(unresolved))
    if unresolved:
        for tok in unresolved:
            print(f"[WARNING] Unresolved placeholder: {tok}")

    # 4c. Optional report
    if report:
        report_path = os.path.splitext(output_path)[0] + ".report.json"
        report_payload = {
            "input": md_path,
            "output": output_path,
            "template": template_path,
            "artifactType": artifact_type,
            "language": resolved_lang,
            "languageFallback": language_fallback,
            "strict": strict,
            "parsedSections": sorted(k for k in sections.keys() if not k.startswith("_")),
            "templatePlaceholders": template_tokens,
            "substitutedPlaceholders": substituted,
            "unresolvedPlaceholders": unresolved,
            "diagrams": sorted(diagrams.keys()),
            "images": sorted(images.keys()),
        }
        try:
            with open(report_path, "w", encoding="utf-8") as rf:
                json.dump(report_payload, rf, indent=2, ensure_ascii=False)
            print(f"[OK] Report file   : {report_path}")
        except Exception as e:
            print(f"[WARNING] Could not write report file: {e}")

    if unresolved and strict:
        print(f"[ERROR] {len(unresolved)} unresolved placeholder(s) and --strict is on. Aborting.")
        sys.exit(1)

    # 5. Save output
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.isdir(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    doc.save(output_path)

    if not os.path.isfile(output_path):
        print("[ERROR] Output file was not created.")
        sys.exit(1)

    # 5b. Patch settings.xml so TOC and fields auto-update on open
    if patch_update_fields(output_path):
        print("[OK] Patched updateFields — TOC will refresh on first open")

    size_kb = round(os.path.getsize(output_path) / 1024, 1)
    print()
    print(f"[OK] Conversion complete!")
    print(f"  File : {output_path}")
    print(f"  Size : {size_kb} KB")
    print(f"  Template applied : YES")
    print()

    # 6. Optionally open
    if open_after:
        print("Opening document...")
        if sys.platform == "win32":
            os.startfile(output_path)
        elif sys.platform == "darwin":
            subprocess.call(["open", output_path])
        else:
            subprocess.call(["xdg-open", output_path])

    print("Conversion FINISHED. Template was applied successfully.")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Convert a Markdown release note to DOCX using a Word template with {{Placeholder}} substitution."
    )
    parser.add_argument(
        "markdown_path",
        nargs="?",
        default=None,
        help="Path to the markdown (.md) file to convert.",
    )
    parser.add_argument(
        "--template",
        dest="template_path",
        default=None,
        help="Path to the .dotx/.docx template file, or the template name/number from the examples folder.",
    )
    parser.add_argument(
        "--output",
        dest="output_path",
        default=None,
        help="Output .docx file path. Default: derived from input path (releasenotesmd→releasenotesdocx).",
    )
    parser.add_argument(
        "--open",
        dest="open_after",
        action="store_true",
        help="Open the generated .docx file after conversion.",
    )
    parser.add_argument(
        "--language",
        dest="language",
        default=None,
        help="BCP-47 language code for Label_* placeholders (en, es, fr, de, it, pt, nl). Falls back to frontmatter 'language:' or --language-fallback.",
    )
    parser.add_argument(
        "--language-fallback",
        dest="language_fallback",
        default="en",
        help="BCP-47 fallback language when a label is missing in the requested --language (default: en).",
    )
    parser.add_argument(
        "--strict",
        dest="strict",
        action="store_true",
        default=True,
        help="Abort with exit code 1 if any {{...}} placeholders remain unresolved (default).",
    )
    parser.add_argument(
        "--no-strict",
        dest="strict",
        action="store_false",
        help="Disable strict mode; unresolved placeholders are warnings only.",
    )
    parser.add_argument(
        "--report",
        dest="report",
        action="store_true",
        default=False,
        help="Write a sibling <output>.report.json listing parsed sections, substituted/unresolved placeholders, diagrams, and images.",
    )
    parser.add_argument(
        "--list-templates",
        dest="list_templates_flag",
        action="store_true",
        help="List all available templates in the examples folder and exit.",
    )

    args = parser.parse_args()

    # --- List templates mode ---
    if args.list_templates_flag:
        templates = list_templates()
        if not templates:
            print("[WARNING] No templates found in the examples folder.")
            print(f"          Expected location: {get_examples_dir()}")
            sys.exit(1)
        print()
        print("Available templates:")
        print()
        for i, t in enumerate(templates, 1):
            size_kb = round(t.stat().st_size / 1024, 1)
            print(f"  {i}. {t.name}  ({size_kb} KB)")
        print()
        print(f"Use --template <number> or --template \"<name>\" to select a template.")
        sys.exit(0)

    # --- Conversion mode: markdown_path is required ---
    if not args.markdown_path:
        parser.error("markdown_path is required for conversion (use --list-templates to see available templates).")

    # Resolve template
    template = args.template_path
    if template:
        # Check if user provided a number (template index)
        if template.isdigit():
            templates = list_templates()
            idx = int(template) - 1
            if 0 <= idx < len(templates):
                template = str(templates[idx])
            else:
                print(f"[ERROR] Invalid template number: {template}. Use --list-templates to see available options.")
                sys.exit(1)
        elif not os.path.isfile(template):
            # Check if it matches a template name in the examples folder
            templates = list_templates()
            match = [t for t in templates if t.name == template or t.stem == template]
            if match:
                template = str(match[0])
            else:
                print(f"[ERROR] Template not found: {template}")
                print("        Use --list-templates to see available templates.")
                sys.exit(1)
    else:
        # First try to match a template by:
        #   1. Exact frontmatter `template:` filename (canonical contract).
        #   2. The artifact's templateHint declared in unified-field-map.json.
        template = None
        try:
            with open(args.markdown_path, "r", encoding="utf-8") as _mf:
                _content = _mf.read()
            _registry = load_unified_field_map()
            _fm = parse_yaml_frontmatter(_content)
            _explicit_tpl = (_fm.get("template") or "").strip()
            if _explicit_tpl:
                for t in list_templates():
                    if t.name == _explicit_tpl or t.stem == Path(_explicit_tpl).stem:
                        template = str(t)
                        print(f"  Template auto-selected via frontmatter 'template:': {t.name}")
                        break
            if not template:
                _atype = detect_artifact_type(_fm, args.markdown_path, _registry)
                if _atype and _registry:
                    hint = _registry["artifacts"][_atype].get("templateHint")
                    if hint:
                        for t in list_templates():
                            if t.stem == hint or t.name == hint:
                                template = str(t)
                                print(f"  Template auto-selected via {_atype} hint: {t.name}")
                                break
        except Exception as _e:
            print(f"[DEBUG] Template hint resolution skipped: {_e}")

        if not template:
            template = find_template(os.path.dirname(os.path.abspath(args.markdown_path)))

    if not template:
        # Multiple templates available — list them and ask the user to choose
        templates = list_templates()
        if templates:
            print()
            print("Multiple templates found. Please select one:")
            print()
            for i, t in enumerate(templates, 1):
                size_kb = round(t.stat().st_size / 1024, 1)
                print(f"  {i}. {t.name}  ({size_kb} KB)")
            print()
            print("Use --template <number> or --template \"<name>\" to select a template.")
            sys.exit(1)
        else:
            print(f"[ERROR] No templates found in the examples folder: {get_examples_dir()}")
            sys.exit(1)

    # Resolve output
    output = args.output_path
    if not output:
        md_dir = os.path.dirname(os.path.abspath(args.markdown_path))
        base = os.path.splitext(os.path.basename(args.markdown_path))[0]
        if md_dir.endswith("releasenotesmd"):
            out_dir = md_dir.replace("releasenotesmd", "releasenotesdocx")
        else:
            out_dir = md_dir
        output = os.path.join(out_dir, base + ".docx")

    convert(
        args.markdown_path,
        template,
        output,
        args.open_after,
        language=args.language,
        language_fallback=args.language_fallback,
        strict=args.strict,
        report=args.report,
    )


if __name__ == "__main__":
    main()
